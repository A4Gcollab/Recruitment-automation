"""Generate A4G_Recruitment_Requirements_v2.2.docx — tight one-pager."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pBdr.append(bottom)
    pPr.append(pBdr)


def tighten(paragraph, before=0, after=2, line=1.05):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def make_run(p, text, *, bold=False, size=10, color=None, italic=False):
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return r


def heading(doc, text):
    p = doc.add_paragraph()
    tighten(p, before=5, after=2)
    make_run(p, text.upper(), bold=True, size=10.5, color=RGBColor(0x1F, 0x3A, 0x5F))
    add_horizontal_line(p)
    return p


def body(doc, text, *, bold=False, size=9.5):
    p = doc.add_paragraph()
    tighten(p, before=0, after=2)
    make_run(p, text, bold=bold, size=size)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    tighten(p, before=0, after=1)
    p.paragraph_format.left_indent = Inches(0.2)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9.5)
    if not p.runs:
        make_run(p, text, size=9.5)
    else:
        p.runs[0].text = text
    return p


def numbered(doc, n, text):
    p = doc.add_paragraph()
    tighten(p, before=0, after=1)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    make_run(p, f"{n}. ", bold=True, size=9.5)
    make_run(p, text, size=9.5)
    return p


def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9.5)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.05

    # Title
    p = doc.add_paragraph()
    tighten(p, before=0, after=0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    make_run(p, "A4G Recruitment Automation — Requirements v2.2",
             bold=True, size=15, color=RGBColor(0x1F, 0x3A, 0x5F))

    # Metadata
    p = doc.add_paragraph()
    tighten(p, before=0, after=4)
    make_run(p,
             "Owner: Sneha   •   Source: Stakeholder meeting 2026-05-18   •   "
             "Updated 2026-05-20   •   Supersedes PRD v2.1 §6.3",
             size=8.5, color=RGBColor(0x55, 0x55, 0x55))
    add_horizontal_line(p)

    # Purpose
    heading(doc, "Purpose")
    body(doc,
         "A zero-cost web tool that helps HR screen, evaluate, and shortlist candidates "
         "~50–70% faster by automating bulk email outreach and the data round-trip with ChatGPT. "
         "The tool does not evaluate candidates itself — ChatGPT does, manually.")

    # Workflow
    heading(doc, "Workflow")
    steps = [
        "HR exports LinkedIn Easy Apply applicants via the free ApplicantSync Chrome extension; "
        "ApplicantSync asynchronously enriches each row with name, email, phone, current title and company, "
        "school, location, LinkedIn URL, resume URL, and an auto-score (~95%+ enrichment in a few hours).",
        "HR imports the ApplicantSync CSV / Sheet into the tool.",
        "Tool exports the candidate list as an XLSX (Excel) for upload to ChatGPT.",
        "ChatGPT fills the verdict and reason columns; HR re-imports the XLSX.",
        "Tool sends the Google Form email in bulk to all selected good-fit candidates.",
        "Tool auto-sends a reminder after 3 days to non-responders (one reminder only).",
        "Tool pulls Google Form responses; exports XLSX with responses + the prior verdict and reason + "
        "empty interview_verdict and reason columns.",
        "ChatGPT fills the interview verdict; HR re-imports the XLSX.",
        "Tool sends the interview-link email in bulk to all selected call-for-interview candidates.",
    ]
    for i, s in enumerate(steps, 1):
        numbered(doc, i, s)

    # In Scope
    heading(doc, "In Scope (v0.2)")
    items = [
        "Import ApplicantSync CSV / Sheet, capturing all fields: name, email, phone, current title, "
        "current company, school, location, LinkedIn URL, applied date, resume URL, ApplicantSync score, "
        "plus any LinkedIn screening-question answers in a JSONB bag.",
        "Export candidate list as XLSX — all captured fields, plus empty verdict and reason columns "
        "for ChatGPT to fill (resume URL included so ChatGPT can optionally read the PDF).",
        "Import evaluated XLSX — match by email, write verdict and reason onto each candidate.",
        "Bulk send Google Form email — HRBP-customizable template, sent only to selected candidates.",
        "Auto-reminder every 3 days if no form response — one reminder per candidate, then stop.",
        "Pull Google Form responses into candidate records (match by email).",
        "Export form responses as XLSX — responses + prior verdict and reason (context) + empty "
        "interview_verdict and reason columns.",
        "Import interview-verdict XLSX — write interview verdict and reason onto each candidate.",
        "Bulk send interview-link email — HRBP-customizable template with free-text note and Zoom link.",
        "Per-campaign customizable templates for Stage-1, reminder, and interview-link emails.",
    ]
    for s in items:
        bullet(doc, s)

    # Out of Scope
    heading(doc, "Out of Scope")
    out = [
        "AI evaluation inside the tool — ChatGPT stays manual and browser-only (no API).",
        "LinkedIn DM outreach — all outreach is email-only.",
        "JD creation automation — HRBPs continue to write JDs manually.",
        "Gmail reply auto-detection — deferred.",
        "Tracker sheet auto-sync back to Google Sheets — deferred.",
    ]
    for s in out:
        bullet(doc, s)

    # Defaults & Constraints
    heading(doc, "Defaults & Constraints")
    cons = [
        "Candidate data source: ApplicantSync Chrome extension, free tier (async enrichment, ~95%+ coverage).",
        "Import / export format: XLSX (Excel).",
        "Reminder cadence: every 3 days, maximum 1 reminder per candidate.",
        "Send window: 9 AM – 8 PM IST, 20 emails per hour, 30–60s gap between sends.",
        "Cost: ₹0 — Gmail SMTP, Google Sheets, ApplicantSync free tier, ChatGPT manual session.",
        "Kill switch: a single environment variable halts all sends instantly.",
    ]
    for s in cons:
        bullet(doc, s)

    out_path = "/mnt/d/workspace/linkedin-automation/A4G_Recruitment_Requirements_v2.2.docx"
    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    build()
